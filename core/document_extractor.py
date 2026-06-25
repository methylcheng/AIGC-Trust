"""
文档文本提取工具
支持 Word (.docx) 和 PDF (.pdf) 格式
"""
import os


def extract_text_from_docx(file_path: str) -> str:
    """
    从 Word 文档提取文本
    
    Args:
        file_path: .docx 文件路径
        
    Returns:
        提取的文本内容
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        # 提取所有段落文本
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        
        return '\n\n'.join(paragraphs)
        
    except Exception as e:
        raise Exception(f"Word文档解析失败: {str(e)}")


def extract_text_from_pdf(file_path: str) -> str:
    """
    从 PDF 文档提取文本
    
    Args:
        file_path: .pdf 文件路径
        
    Returns:
        提取的文本内容
    """
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        # 提取每一页的文本
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text.strip())
        
        if not text_parts:
            raise Exception("PDF中未找到可提取的文本（可能是扫描版或图片）")
        
        return '\n\n'.join(text_parts)
        
    except Exception as e:
        raise Exception(f"PDF文档解析失败: {str(e)}")


def extract_text_from_file(file_path: str) -> str:
    """
    根据文件扩展名自动选择提取方法
    
    Args:
        file_path: 文件路径
        
    Returns:
        提取的文本内容
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.pdf':
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，仅支持 .txt, .docx, .pdf")


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        try:
            text = extract_text_from_file(file_path)
            print(f"提取成功！共 {len(text)} 字符")
            print(f"\n前500字符预览:\n{text[:500]}...")
        except Exception as e:
            print(f"提取失败: {e}")
    else:
        print("用法: python document_extractor.py <文件路径>")
